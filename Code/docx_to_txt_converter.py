# DOCX to TXT Converter in Jupyter Notebook

# Method 1: Using python-docx (Recommended)
# First install the required library:
# !pip install python-docx

from docx import Document
import os

def docx_to_txt_method1(docx_path, txt_path=None):
    """
    Convert DOCX to TXT using python-docx library
    
    Args:
        docx_path (str): Path to the input DOCX file
        txt_path (str): Path for output TXT file (optional)
    
    Returns:
        str: Extracted text content
    """
    # Load the document
    doc = Document(docx_path)
    
    # Extract text from all paragraphs
    text_content = []
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)
    
    # Join paragraphs with newlines
    full_text = '\n'.join(text_content)
    
    # Save to file if path provided
    if txt_path:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        print(f"Text saved to: {txt_path}")
    
    return full_text

# Method 2: Using docx2txt (Alternative)
# First install: !pip install docx2txt

import docx2txt

def docx_to_txt_method2(docx_path, txt_path=None):
    """
    Convert DOCX to TXT using docx2txt library
    
    Args:
        docx_path (str): Path to the input DOCX file
        txt_path (str): Path for output TXT file (optional)
    
    Returns:
        str: Extracted text content
    """
    # Extract text directly
    text = docx2txt.process(docx_path)
    
    # Save to file if path provided
    if txt_path:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Text saved to: {txt_path}")
    
    return text

# Method 3: Batch conversion for multiple files
def batch_docx_to_txt(input_folder, output_folder=None, method='python-docx'):
    """
    Convert multiple DOCX files to TXT
    
    Args:
        input_folder (str): Folder containing DOCX files
        output_folder (str): Folder for TXT files (optional)
        method (str): 'python-docx' or 'docx2txt'
    """
    if output_folder is None:
        output_folder = input_folder
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all DOCX files
    for filename in os.listdir(input_folder):
        if filename.lower().endswith('.docx'):
            docx_path = os.path.join(input_folder, filename)
            txt_filename = filename.replace('.docx', '.txt')
            txt_path = os.path.join(output_folder, txt_filename)
            
            try:
                if method == 'python-docx':
                    docx_to_txt_method1(docx_path, txt_path)
                else:
                    docx_to_txt_method2(docx_path, txt_path)
                print(f"✓ Converted: {filename}")
            except Exception as e:
                print(f"✗ Error converting {filename}: {str(e)}")

# Example usage:

# Single file conversion
if __name__ == "__main__":
    # Example 1: Convert single file using method 1
    docx_file = "example.docx"  # Replace with your file path
    txt_file = "output.txt"     # Output file path
    
    # Check if file exists before processing
    if os.path.exists(docx_file):
        # Method 1: python-docx
        text_content = docx_to_txt_method1(docx_file, txt_file)
        print("Preview of extracted text:")
        print(text_content[:500] + "..." if len(text_content) > 500 else text_content)
        
        # Method 2: docx2txt (uncomment to use)
        # text_content = docx_to_txt_method2(docx_file, "output_method2.txt")
    else:
        print(f"File not found: {docx_file}")
    
    # Example 2: Batch conversion
    input_dir = "docx_files"    # Replace with your input folder
    output_dir = "txt_files"    # Replace with your output folder
    
    if os.path.exists(input_dir):
        print("\nStarting batch conversion...")
        batch_docx_to_txt(input_dir, output_dir, method='python-docx')
        print("Batch conversion completed!")
    else:
        print(f"Input directory not found: {input_dir}")

# Advanced example: Extract with formatting preservation
def docx_to_txt_advanced(docx_path, txt_path=None, preserve_formatting=True):
    """
    Advanced DOCX to TXT conversion with formatting options
    
    Args:
        docx_path (str): Path to input DOCX file
        txt_path (str): Path for output TXT file
        preserve_formatting (bool): Whether to preserve basic formatting
    
    Returns:
        str: Extracted text content
    """
    doc = Document(docx_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if preserve_formatting:
            # Add extra spacing for headings (based on style)
            if paragraph.style.name.startswith('Heading'):
                text_parts.append(f"\n{'='*50}")
                text_parts.append(paragraph.text.upper())
                text_parts.append('='*50)
            else:
                text_parts.append(paragraph.text)
        else:
            text_parts.append(paragraph.text)
    
    # Process tables if any
    for table in doc.tables:
        text_parts.append("\n[TABLE]")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            text_parts.append(" | ".join(row_text))
        text_parts.append("[/TABLE]\n")
    
    full_text = '\n'.join(text_parts)
    
    if txt_path:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        print(f"Advanced conversion saved to: {txt_path}")
    
    return full_text

# Usage in Jupyter Notebook:
"""
# In your Jupyter notebook cell:

# Install required libraries
!pip install python-docx docx2txt

# Import and run conversion
from docx import Document
import docx2txt

# Convert single file
text = docx_to_txt_method1('your_file.docx', 'output.txt')
print(text[:200])  # Preview first 200 characters

# Or for batch conversion
batch_docx_to_txt('input_folder/', 'output_folder/')
"""